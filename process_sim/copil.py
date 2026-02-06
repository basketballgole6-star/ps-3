# collect_py_to_docx.py
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# Directories to skip
SKIP_DIRS = {'.git', '.venv', 'venv', '__pycache__', '_pycache_'}

def set_run_monospace(run, font_name='Courier New', size_pt=9):
    try:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        rPr = run._element.rPr
        rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass

def add_code_block(doc, code_text):
    if code_text == '':
        doc.add_paragraph('')
        return
    for line in code_text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_monospace(run)

def collect_py_files_to_docx(root_dir='.', out_docx='collected_code.docx'):
    doc = Document()
    doc.core_properties.title = 'Collected Python Files'
    doc.core_properties.subject = f'Collected from {os.path.abspath(root_dir)}'

    for dirpath, dirnames, filenames in os.walk(root_dir):
        # Filter out skip dirs and sort for reproducibility
        dirnames[:] = sorted(d for d in dirnames if d not in SKIP_DIRS)
        rel_dir = os.path.relpath(dirpath, root_dir)
        py_files = sorted(f for f in filenames if f.endswith('.py'))
        if not py_files:
            continue

        folder_heading = rel_dir if rel_dir != '.' else os.path.basename(os.path.abspath(root_dir))
        doc.add_heading(f'Folder: {folder_heading}', level=2)

        for fname in py_files:
            file_path = os.path.join(dirpath, fname)
            doc.add_heading(f'File: {fname}', level=3)
            doc.add_paragraph(f'Path: {os.path.relpath(file_path, root_dir)}')

            try:
                with open(file_path, 'r', encoding='utf-8') as fh:
                    content = fh.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as fh:
                    content = fh.read()

            add_code_block(doc, content)
            doc.add_paragraph('')

    doc.save(out_docx)
    print(f'Saved {out_docx}')

if __name__ == '__main__':
    collect_py_files_to_docx(root_dir='.', out_docx='collected_code.docx')
